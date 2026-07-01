#!/usr/bin/env python3
"""生成 kylin-doctor 项目架构分析与作品介绍 Word 文档"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading_elm = cell._element.get_or_add_tcPr()
    shading = shading_elm.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): color
    })
    shading_elm.append(shading)

def add_heading_with_style(doc, text, level=1):
    """添加带样式的标题"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x1e, 0x29, 0x3b)
    return heading

def add_table_with_style(doc, headers, rows, col_widths=None):
    """添加带样式的表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, '1e293b')
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
                run.font.bold = True
                run.font.size = Pt(10)

    # 数据行
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
            if i % 2 == 0:
                set_cell_shading(cell, 'f8fafc')

    return table

def main():
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft YaHei'
    font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    # ==================== 封面 ====================
    doc.add_paragraph()
    doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('🔍 kylin-doctor')
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x06, 0xb6, 0xd4)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('银河麒麟桌面系统自我诊断工具')
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)

    doc.add_paragraph()

    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = desc.add_run('架构分析与作品介绍')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)

    doc.add_paragraph()
    doc.add_paragraph()

    version_info = doc.add_paragraph()
    version_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = version_info.add_run('版本: v0.3.5 | 语言: Rust | 协议: MIT')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)

    doc.add_page_break()

    # ==================== 目录 ====================
    add_heading_with_style(doc, '📋 目录', level=1)

    toc_items = [
        '1. 项目简介 — 这是什么？',
        '2. 为什么要做这个？',
        '3. 核心功能一览',
        '4. 系统架构设计',
        '5. 五大检测模块详解',
        '6. AI 智能问答系统',
        '7. 技术亮点',
        '8. 版本历史',
        '9. 未来规划',
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(4)

    doc.add_page_break()

    # ==================== 1. 项目简介 ====================
    add_heading_with_style(doc, '1. 项目简介 — 这是什么？', level=1)

    doc.add_paragraph(
        'kylin-doctor（麒麟医生）是一款专为银河麒麟桌面系统设计的自我诊断工具。'
        '它就像一位"系统医生"，能够自动检查电脑的健康状况，发现问题并给出修复建议。'
    )

    doc.add_paragraph()
    add_heading_with_style(doc, '一句话总结', level=2)
    p = doc.add_paragraph()
    run = p.add_run('让 Linux 系统维护变得像体检一样简单 —— 一键扫描，AI 诊断，自动修复。')
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x06, 0xb6, 0xd4)

    doc.add_paragraph()
    add_heading_with_style(doc, '关键数据', level=2)

    add_table_with_style(doc,
        ['指标', '数值'],
        [
            ['总代码量', '10,907 行'],
            ['Rust 代码', '10,907 行'],
            ['检测模块', '5 大模块，49 项检查'],
            ['支持的 AI 模型', '3 种（Ollama/OpenAI/Anthropic）'],
            ['支持架构', 'amd64 + arm64（均为 musl 静态编译）'],
            ['当前版本', 'v0.3.5'],
        ]
    )

    doc.add_page_break()

    # ==================== 2. 为什么要做这个？ ====================
    add_heading_with_style(doc, '2. 为什么要做这个？', level=1)

    doc.add_paragraph('在日常工作中，我们经常遇到以下场景：')

    scenarios = [
        ('😵 用户报障', '电脑卡了、连不上网、软件打不开，运维人员需要逐个排查，耗时耗力'),
        ('🔧 系统维护', '定期巡检需要检查几十项指标，人工操作容易遗漏'),
        ('📚 技术门槛', '很多问题需要深厚的 Linux 知识才能诊断，普通用户束手无策'),
        ('🏭 工控场景', '银河麒麟系统广泛应用于工控机、嵌入式设备，远程维护困难'),
    ]

    for title, desc in scenarios:
        p = doc.add_paragraph()
        run = p.add_run(title + '：')
        run.font.bold = True
        p.add_run(desc)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('kylin-doctor 就是为了解决这些问题而生的。')
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x06, 0xb6, 0xd4)

    doc.add_page_break()

    # ==================== 3. 核心功能一览 ====================
    add_heading_with_style(doc, '3. 核心功能一览', level=1)

    add_heading_with_style(doc, '三种使用方式', level=2)

    add_table_with_style(doc,
        ['方式', '命令', '适用场景'],
        [
            ['命令行 (CLI)', 'kylin-doctor scan', '远程 SSH 维护、脚本集成'],
            ['Web 仪表盘', 'kylin-doctor-web', '可视化监控、实时诊断'],
            ['AI 对话', 'kylin-doctor chat', '自然语言提问、智能诊断'],
        ]
    )

    doc.add_paragraph()
    add_heading_with_style(doc, '六大核心能力', level=2)

    features = [
        ('🔍 全面扫描', '5 大模块、49 项检查，覆盖系统、硬件、软件、安全、性能'),
        ('🤖 AI 诊断', '集成大语言模型，用自然语言描述问题，AI 分析并给出解决方案'),
        ('🛠️ 自动修复', '发现的问题可以一键修复，也可以先预览再决定'),
        ('📊 报告导出', '支持 JSON 和 HTML 格式导出诊断报告'),
        ('🌐 双端支持', 'CLI 命令行 + Web 仪表盘，满足不同场景需求'),
        ('📦 即装即用', 'deb 包一键安装，静态编译无依赖'),
    ]

    for title, desc in features:
        p = doc.add_paragraph()
        run = p.add_run(title + '：')
        run.font.bold = True
        p.add_run(desc)

    doc.add_page_break()

    # ==================== 4. 系统架构设计 ====================
    add_heading_with_style(doc, '4. 系统架构设计', level=1)

    add_heading_with_style(doc, '整体架构图', level=2)

    # 用文字描述架构图
    arch_text = '''
┌─────────────────────────────────────────────────────────────┐
│                      用户界面层                              │
│  ┌──────────────────┐  ┌──────────────────────────────────┐ │
│  │   CLI 命令行      │  │   Web 仪表盘 (单文件 HTML)       │ │
│  │   scan/chat/fix  │  │   ECharts + WebSocket + AI Chat  │ │
│  └────────┬─────────┘  └───────────────┬──────────────────┘ │
│           │                            │                    │
├───────────┼────────────────────────────┼────────────────────┤
│           │      核心逻辑层            │                    │
│           │   ┌────────────────────────┴──────────────┐    │
│           └──→│         kylin-doctor-core              │    │
│               │  ┌─────────┐ ┌─────────┐ ┌──────────┐ │    │
│               │  │ 检测器   │ │ LLM 引擎│ │ 知识库   │ │    │
│               │  │ 5 模块   │ │ 3 Provider│ │ RAG     │ │    │
│               │  └─────────┘ └─────────┘ └──────────┘ │    │
│               └───────────────────────────────────────┘    │
│                                                            │
├────────────────────────────────────────────────────────────┤
│                      系统层                                │
│    /proc/stat  /proc/meminfo  /proc/diskstats  systemd     │
│    /sys/class  dmesg  smartctl  apt  snap  flatpak         │
└────────────────────────────────────────────────────────────┘
'''
    p = doc.add_paragraph()
    run = p.add_run(arch_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

    doc.add_paragraph()
    add_heading_with_style(doc, '三个 Crate 的分工', level=2)

    add_table_with_style(doc,
        ['Crate', '代码量', '职责', '二进制'],
        [
            ['kylin-doctor-core', '7,310 行', '核心检测逻辑 + LLM 集成 + 知识库', '库（被其他两个依赖）'],
            ['kylin-doctor-cli', '2,037 行', '命令行界面 + 交互式聊天', 'kylin-doctor'],
            ['kylin-doctor-web', '1,281 行', 'Web 服务器 + API + 前端', 'kylin-doctor-web'],
        ]
    )

    doc.add_page_break()

    # ==================== 5. 五大检测模块详解 ====================
    add_heading_with_style(doc, '5. 五大检测模块详解', level=1)

    modules = [
        {
            'name': '💻 系统模块 (system)',
            'count': '5 项检查',
            'checks': [
                ('磁盘空间', '检查各挂载点使用率，超过 80% 告警，超过 90% 严重'),
                ('失败服务', '扫描 systemd 失败的服务，可自动重启'),
                ('僵尸进程', '检测 Z 状态进程，避免资源泄漏'),
                ('内核错误', '分析 dmesg 日志，发现硬件/驱动错误'),
                ('系统负载', '对比 1/5/15 分钟负载与 CPU 核心数'),
            ]
        },
        {
            'name': '🔧 硬件模块 (hardware)',
            'count': '10 项检查',
            'checks': [
                ('CPU 温度', '读取 thermal_zone，超过 80°C 告警'),
                ('GPU 状态', '检测 NVIDIA 显卡温度和显存使用'),
                ('内存使用', '监控可用内存，超过 85% 告警'),
                ('磁盘健康', '通过 SMART 检测磁盘健康状态'),
                ('磁盘寿命', '监控 SSD 磨损等级和重映射扇区'),
                ('磁盘速度', '采样计算实际读写速度'),
                ('网卡状态', '检测链路状态和错误包率'),
                ('USB 设备', '统计外设数量，检查打印/扫描服务'),
                ('主板信息', '读取 BIOS 版本和 CMOS 电池电压'),
                ('I/O 错误', '监控磁盘 I/O 错误计数'),
            ]
        },
        {
            'name': '📦 软件模块 (software)',
            'count': '10 项检查',
            'checks': [
                ('APT 状态', '检测损坏包、可升级包数量'),
                ('软件源', '检查源配置和 GPG 密钥'),
                ('依赖冲突', '分析依赖关系，检测冲突'),
                ('运行环境', '检查 Python、pip 等运行时'),
                ('中文字体', '检测中文字体安装情况'),
                ('字体渲染', '检查 fontconfig CJK 配置'),
                ('Wine', '检测 Wine 安装和前缀状态'),
                ('安卓兼容', '检查 Anbox/Waydroid 状态'),
                ('通用包', '检查 Snap/Flatpak 健康状态'),
                ('应用审计', '检查应用权限和更新状态'),
            ]
        },
        {
            'name': '🔒 安全模块 (security)',
            'count': '12 项检查',
            'checks': [
                ('空密码', '扫描 /etc/shadow 空密码账户'),
                ('特权账户', '检测非 root 的 UID 0 账户'),
                ('过期账户', '检查密码过期和长期未登录账户'),
                ('SUID 文件', '扫描可疑的 SUID 可执行文件'),
                ('目录权限', '检查关键目录权限是否合规'),
                ('SSH 配置', '检查 root 登录、密码认证等设置'),
                ('防火墙', '检查 ufw/iptables 状态和规则'),
                ('开放端口', '扫描高风险端口（telnet/ftp/mysql 等）'),
                ('登录失败', '分析暴力破解尝试，可安装 fail2ban'),
                ('审计日志', '检查 auditd 状态和 sudo 使用情况'),
                ('已知漏洞', '检查内核版本和待安装安全更新'),
                ('密码策略', '检查密码有效期和最小长度'),
            ]
        },
        {
            'name': '⚡ 性能模块 (performance)',
            'count': '12 项检查',
            'checks': [
                ('CPU 使用率', '实时采样 /proc/stat，与 top 一致'),
                ('CPU 调度', '分析调度延迟，检测 CPU 争用'),
                ('负载趋势', '对比 1/5/15 分钟负载变化'),
                ('内存性能', '监控 Swap、脏页、透明大页'),
                ('内存碎片', '分析 buddyinfo 高阶空闲块比例'),
                ('磁盘 I/O 延迟', '采样计算平均 I/O 延迟'),
                ('磁盘 IOPS', '统计每秒读写操作数和队列深度'),
                ('网络连接', '统计 TCP 连接数和 TIME_WAIT'),
                ('网络延迟', 'ping 默认网关测量延迟'),
                ('网络带宽', '检测链路速度、错误率、累计吞吐'),
                ('桌面合成器', '检测 KWin/Mutter 等资源占用'),
                ('电池/电源', '监控电池健康和 CPU 调频策略'),
            ]
        },
    ]

    for mod in modules:
        add_heading_with_style(doc, mod['name'], level=2)
        p = doc.add_paragraph()
        run = p.add_run(f'共 {mod["count"]}')
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)

        for check_name, check_desc in mod['checks']:
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(check_name + '：')
            run.font.bold = True
            p.add_run(check_desc)

    doc.add_page_break()

    # ==================== 6. AI 智能问答系统 ====================
    add_heading_with_style(doc, '6. AI 智能问答系统', level=1)

    doc.add_paragraph(
        'kylin-doctor 集成了大语言模型（LLM），用户可以用自然语言提问，'
        'AI 会自动调用检测工具分析系统状态，然后给出通俗易懂的诊断结果和修复建议。'
    )

    doc.add_paragraph()
    add_heading_with_style(doc, '支持的 AI 模型', level=2)

    add_table_with_style(doc,
        ['模型', '类型', '特点', '适用场景'],
        [
            ['Ollama (qwen2.5:3b)', '本地部署', '离线可用，隐私安全', '工控机、内网环境'],
            ['通义千问 / DeepSeek', '云端 API', '能力强，响应快', '有网络的办公环境'],
            ['Anthropic Claude', '云端 API', '推理能力强', '复杂问题诊断'],
        ]
    )

    doc.add_paragraph()
    add_heading_with_style(doc, 'AI 对话示例', level=2)

    examples = [
        '用户：电脑好卡啊，怎么回事？',
        'AI：让我帮您检查一下...（自动调用扫描工具）',
        'AI：发现以下问题：',
        '  1. CPU 使用率 92%，有进程占用过高',
        '  2. 内存使用率 88%，可用内存不足',
        '  3. 磁盘 / 使用率 95%，空间即将耗尽',
        '',
        '建议操作：',
        '  • 运行 top 命令查看高耗进程',
        '  • 清理磁盘空间：sudo apt clean',
        '  • 扩容或迁移数据',
    ]

    p = doc.add_paragraph()
    run = p.add_run('\n'.join(examples))
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)

    doc.add_paragraph()
    add_heading_with_style(doc, '技术实现', level=2)

    ai_features = [
        ('Function Calling', 'AI 可以调用预定义的检测工具，获取实时系统数据'),
        ('流式输出', 'AI 回答逐字显示，体验流畅'),
        ('上下文管理', '自动维护对话历史，支持多轮对话'),
        ('混合模式', '本地模型优先，失败自动切换云端'),
        ('知识库增强 (RAG)', '支持导入技术文档，增强 AI 的专业知识'),
    ]

    for title, desc in ai_features:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(title + '：')
        run.font.bold = True
        p.add_run(desc)

    doc.add_page_break()

    # ==================== 7. 技术亮点 ====================
    add_heading_with_style(doc, '7. 技术亮点', level=1)

    highlights = [
        {
            'title': '🦀 Rust 语言开发',
            'desc': '内存安全、零成本抽象、高性能。单二进制部署，无运行时依赖。'
        },
        {
            'title': '📦 静态编译 (musl)',
            'desc': '使用 musl 静态编译，彻底消除 glibc 版本依赖。'
                    'amd64 和 arm64 双架构均为静态链接（arm64 使用 cross 工具进行交叉编译），'
                    '可在任何 Linux 系统运行，包括低版本 glibc 的工控机。'
        },
        {
            'title': '⏱️ 命令超时保护',
            'desc': '所有外部命令（apt、snap、find 等）都有超时保护（默认 10 秒，apt 30 秒），'
                    '防止 apt 锁被占用时扫描卡死。'
        },
        {
            'title': '🌐 WebSocket 实时通信',
            'desc': '扫描进度、AI 对话都通过 WebSocket 实时推送，'
                    '用户体验流畅。'
        },
        {
            'title': '📡 双层 Channel 流式桥接',
            'desc': '使用 std::sync::mpsc + tokio::sync::mpsc 双层通道，'
                    '将同步的 LLM 流式回调桥接到异步 WebSocket 发送。'
        },
        {
            'title': '🧠 RAG 知识库',
            'desc': '支持导入技术文档，通过向量检索增强 AI 的专业知识。'
                    '使用 Ollama 生成嵌入向量，余弦相似度搜索。'
        },
        {
            'title': '🎨 单文件前端',
            'desc': 'Web 仪表盘是单个 HTML 文件（587 行），'
                    'CSS + JS 全部内嵌，零外部依赖。'
                    '暗色主题，ECharts 图表，响应式设计。'
        },
        {
            'title': '🔒 安全检测全面',
            'desc': '12 项安全检查，覆盖空密码、特权账户、SUID 文件、'
                    'SSH 配置、防火墙、开放端口、暴力破解等。'
        },
    ]

    for h in highlights:
        p = doc.add_paragraph()
        run = p.add_run(h['title'])
        run.font.bold = True
        run.font.size = Pt(12)
        doc.add_paragraph(h['desc'])
        doc.add_paragraph()

    doc.add_page_break()

    # ==================== 8. 版本历史 ====================
    add_heading_with_style(doc, '8. 版本历史', level=1)

    add_table_with_style(doc,
        ['版本', '日期', '主要变更'],
        [
            ['v0.3.5', '2026-07-01', 'arm64 musl 静态编译完成（使用 cross 工具），双架构 deb 包均为静态链接'],
            ['v0.3.3', '2026-06-29', '修复 CPU 实时显示、修复 iowait 计算逻辑'],
            ['v0.3.2', '2026-06-27', '修复 CPU 显示 0%、支持 musl 静态编译'],
            ['v0.3.1', '2026-06-23', '配置文件直接写入 API Key、外部命令超时保护'],
            ['v0.3.0', '2026-06-15', 'AI 问答全面升级、Anthropic Claude 支持、deb 打包'],
            ['v0.2.0', '2026-06-12', 'AI 问答优化、安装脚本改进'],
            ['v0.1.0', '2026-06-02', '初始版本'],
        ]
    )

    doc.add_page_break()

    # ==================== 9. 未来规划 ====================
    add_heading_with_style(doc, '9. 未来规划', level=1)

    plans = [
        ('📈 监控模式', '增加守护进程模式，定时采集指标，支持历史趋势分析'),
        ('🔔 告警通知', '支持邮件、钉钉、企业微信等告警通知方式'),
        ('🌐 Web 增强', '增加用户认证、多节点管理、远程执行'),
        ('🤖 AI 增强', '支持更多模型、增加自动修复建议、知识库自动学习'),
        ('📊 可视化', '增加实时监控仪表盘、历史趋势图表'),
        ('🔌 插件系统', '支持用户自定义检测规则和修复脚本'),
    ]

    for title, desc in plans:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(title + '：')
        run.font.bold = True
        p.add_run(desc)

    doc.add_paragraph()
    doc.add_paragraph()

    # 结尾
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('— 感谢关注 kylin-doctor —')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x06, 0xb6, 0xd4)
    run.font.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('GitHub: https://github.com/fanwenzhu/kylin-doctor')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)

    # 保存
    output_path = '/home/agent/projects/kylin-doctor/docs/kylin-doctor-架构分析与作品介绍.docx'
    doc.save(output_path)
    print(f'✅ 文档已生成: {output_path}')

if __name__ == '__main__':
    main()
