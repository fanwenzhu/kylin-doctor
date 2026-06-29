#!/usr/bin/env python3
"""生成 kylin-doctor v0.2.0 ~ v0.3.1 更新说明 Word 文档"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# ── 样式设置 ──
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.5

# ── 标题 ──
title = doc.add_heading('', level=0)
run = title.add_run('kylin-doctor 更新日志')
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x06, 0xb6, 0xd4)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('v0.2.0 → v0.3.1  |  2026-06-12 ~ 2026-06-24')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)
run.font.italic = True

doc.add_paragraph()

# ── 项目简介 ──
intro = doc.add_paragraph()
run = intro.add_run('kylin-doctor')
run.bold = True
run.font.color.rgb = RGBColor(0x06, 0xb6, 0xd4)
intro.add_run(' 是银河麒麟桌面系统的自我诊断工具，支持系统、硬件、软件、安全、性能五大模块的自动检测，集成 AI 智能问答，提供 CLI 命令行和 Web 仪表盘两种使用方式。')

doc.add_paragraph()

# ── 版本总览表 ──
doc.add_heading('版本总览', level=1)

table = doc.add_table(rows=4, cols=3)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['版本', '日期', '核心亮点']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True

data = [
    ['v0.2.0', '2026-06-12', 'AI 问答流式输出、终端 Markdown 渲染、安装脚本增强'],
    ['v0.3.0', '2026-06-15', 'deb 打包、Anthropic Claude 支持、Web 端全面升级'],
    ['v0.3.1', '2026-06-24', '配置文件 API Key、扫描超时保护、浏览器兼容性修复'],
]
for row_idx, row_data in enumerate(data):
    for col_idx, val in enumerate(row_data):
        table.rows[row_idx + 1].cells[col_idx].text = val

doc.add_paragraph()

# ── v0.3.1 ──
doc.add_heading('v0.3.1 — 配置优化与稳定性修复', level=1)
doc.add_paragraph('发布日期：2026-06-24')

doc.add_heading('新增功能', level=2)
items = [
    ('配置文件直接写入 API Key', 'config.toml 新增 api_key 字段，无需再设置环境变量。优先级：api_key（配置文件）> api_key_env（环境变量），适用于所有云端供应商（通义千问/DeepSeek/月之暗面/Anthropic/自定义）。'),
]
for title_text, desc in items:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(title_text)
    run.bold = True
    p.add_run(f'：{desc}')

doc.add_heading('问题修复', level=2)
items = [
    ('Web 仪表盘扫描按钮无响应', '移除 JavaScript 中的正则 lookbehind 断言，兼容工控机等旧版浏览器（ES2018 以下），修复 SyntaxError 导致整个 script 块不执行的问题。'),
    ('Web 仪表盘扫描挂起', '所有外部系统命令（apt-get、dpkg、snap、flatpak 等）增加超时保护，防止 apt 锁被占用时扫描请求永久阻塞。新增 command_output_with_timeout() 工具函数，子进程超时自动 kill。'),
]
for title_text, desc in items:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(title_text)
    run.bold = True
    p.add_run(f'：{desc}')

doc.add_paragraph()

# ── v0.3.0 ──
doc.add_heading('v0.3.0 — AI 问答全面升级 + Anthropic Claude 支持', level=1)
doc.add_paragraph('发布日期：2026-06-15')

doc.add_heading('新增功能', level=2)
items = [
    ('deb 安装包支持', '新增 build-deb.sh 一键打包脚本，支持 amd64/arm64 架构及交叉编译，发布 GitHub Release 可直接下载安装。'),
    ('Anthropic Claude 支持', '新增 AnthropicProvider，支持 Anthropic Messages API 原生协议，流式输出 + Function Calling 全支持。'),
    ('Web 端流式输出', 'WebSocket 支持 stream_start/stream_chunk/stream_end 消息类型，AI 回复逐 token 实时显示，打字光标动画。'),
    ('Web 端 Markdown 渲染', '内联 renderMd() 解析器，支持代码块、粗体、斜体、列表、标题、链接、表格、引用。'),
    ('Web 端工具调用可视化', '黄色 spinner 显示工具执行状态，扫描结果支持折叠/展开。'),
    ('CLI 上下文管理', 'clear/清屏/重置 重置对话，history/历史 查看对话摘要。'),
    ('CLI 错误恢复', '流式输出失败时自动回退到非流式批量输出，保证对话不中断。'),
]
for title_text, desc in items:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(title_text)
    run.bold = True
    p.add_run(f'：{desc}')

doc.add_heading('改进', level=2)
items = [
    'CLI 流式输出全覆盖：普通回复、工具调用后回复、hybrid 回退路径全部走流式输出。',
    'CLI 统一输出格式：所有回复路径使用 🤖 助手: 前缀 + 空行分隔，视觉体验一致。',
    'Web 端后端架构优化：使用 tokio::sync::mpsc + std::sync::mpsc 双层 channel 桥接同步回调与异步 WebSocket。',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('问题修复', level=2)
items = [
    'deb 包依赖问题：移除 procps/coreutils 硬依赖，改为 Recommends，缺失不崩溃。',
    'CLI 消息恢复错误：修复工具调用流式失败时消息丢失的问题。',
    'Web 端有序列表渲染：修复有序列表被错误包裹在 <ul> 中的 bug。',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph()

# ── v0.2.0 ──
doc.add_heading('v0.2.0 — AI 问答体验优化', level=1)
doc.add_paragraph('发布日期：2026-06-12')

doc.add_heading('新增功能', level=2)
items = [
    ('AI 问答流式输出', '支持 SSE 流式响应，逐字显示 AI 回复，告别等待。'),
    ('终端 Markdown 渲染', '支持代码块、粗体、斜体、列表、标题、链接，终端输出更美观。'),
    ('安装脚本日志系统', '安装过程记录到 /var/log/kylin-doctor-install.log，出错时可追溯。'),
    ('依赖冲突自动修复', '--fix-deps 选项自动处理 libssl-dev 版本冲突。'),
    ('默认配置文件', '安装时自动创建 ~/.kylin-doctor/config.toml 并提供详细注释。'),
]
for title_text, desc in items:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(title_text)
    run.bold = True
    p.add_run(f'：{desc}')

doc.add_heading('改进', level=2)
items = [
    'AI 问答体验：移除调试输出，使用 ✅/❌ 图标表示操作状态。',
    '默认模型从 qwen2.5:1.5b 改为 qwen2.5:3b，平衡速度和质量。',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('问题修复', level=2)
items = [
    '安装脚本静默退出：修复 curl 下载失败时脚本无提示退出的问题。',
    '错误处理机制：添加 ERR trap 显示具体错误行号、命令和解决方案。',
    '网络超时处理：curl 添加 --connect-timeout 和 --retry 参数。',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph()

# ── 安装方式 ──
doc.add_heading('安装与升级', level=1)

doc.add_heading('全新安装', level=2)
p = doc.add_paragraph()
p.add_run('一键安装脚本：').bold = True
doc.add_paragraph('bash <(curl -sSL https://raw.githubusercontent.com/fanwenzhu/kylin-doctor/master/install.sh)', style='List Bullet')

p = doc.add_paragraph()
p.add_run('deb 包安装（x86_64）：').bold = True
doc.add_paragraph('sudo dpkg -i kylin-doctor_0.3.1_amd64.deb', style='List Bullet')

doc.add_heading('升级已有安装', level=2)
p = doc.add_paragraph()
p.add_run('源码编译升级（推荐，适用于所有架构）：').bold = True
doc.add_paragraph('cd kylin-doctor && git pull && sudo ./install.sh --upgrade', style='List Bullet')

p = doc.add_paragraph()
p.add_run('deb 包覆盖安装（仅 x86_64）：').bold = True
doc.add_paragraph('sudo dpkg -i kylin-doctor_0.3.1_amd64.deb', style='List Bullet')

doc.add_paragraph()

# ── 下载地址 ──
doc.add_heading('下载地址', level=1)
p = doc.add_paragraph()
p.add_run('GitHub Release: ').bold = True
p.add_run('https://github.com/fanwenzhu/kylin-doctor/releases/tag/v0.3.1')

p = doc.add_paragraph()
p.add_run('项目主页: ').bold = True
p.add_run('https://github.com/fanwenzhu/kylin-doctor')

doc.add_paragraph()

# ── 页脚 ──
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run('kylin-doctor — 银河麒麟桌面系统自我诊断工具')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)

# 保存
output_path = '/home/agent/projects/kylin-doctor/docs/kylin-doctor-update-v0.2.0-v0.3.1.docx'
doc.save(output_path)
print(f'✅ 已生成: {output_path}')
