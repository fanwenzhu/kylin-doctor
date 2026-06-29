#!/usr/bin/env python3
"""生成 kylin-doctor 项目介绍 Word 文档"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# 设置默认字体
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)

# 标题
title = doc.add_heading('🏥 kylin-doctor', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('银河麒麟桌面系统自我诊断工具')
run.font.size = Pt(16)
run.font.bold = True
run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

doc.add_paragraph()

# 项目目标
doc.add_heading('🎯 项目目标', level=1)
doc.add_paragraph('解决银河麒麟桌面系统运维中的痛点：')
doc.add_paragraph('系统出问题时，不知道从哪查起', style='List Bullet')
doc.add_paragraph('运维人员需要掌握大量 Linux 命令', style='List Bullet')
doc.add_paragraph('工控机现场排查效率低', style='List Bullet')

p = doc.add_paragraph()
run = p.add_run('一句话定位：')
run.bold = True
p.add_run('让系统诊断像"看病"一样简单 —— 扫描、诊断、开处方、自动修复。')

# 核心功能
doc.add_heading('✨ 核心功能', level=1)

doc.add_heading('1. 五维度全面诊断', level=2)

# 创建表格
table = doc.add_table(rows=6, cols=2)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['维度', '检测内容']
for i, header in enumerate(headers):
    table.rows[0].cells[i].text = header
    for paragraph in table.rows[0].cells[i].paragraphs:
        for run in paragraph.runs:
            run.bold = True

data = [
    ('🔧 硬件', '温度、内存、磁盘健康、GPU、网卡'),
    ('💻 系统', '磁盘空间、服务状态、僵尸进程、内核日志'),
    ('📦 软件', '包管理、字体、运行时、兼容层'),
    ('🔒 安全', '密码策略、SSH配置、防火墙、漏洞扫描'),
    ('⚡ 性能', 'CPU/内存/磁盘IO/网络/桌面合成器分析'),
]

for i, (dim, content) in enumerate(data, 1):
    table.rows[i].cells[0].text = dim
    table.rows[i].cells[1].text = content

doc.add_paragraph()

doc.add_heading('2. AI 智能助手', level=2)
doc.add_paragraph('接入本地大模型（Ollama + Qwen2.5）或云端模型（通义千问/DeepSeek/Claude）', style='List Bullet')
doc.add_paragraph('支持自然语言问答："我的系统为什么变慢了？"', style='List Bullet')
p = doc.add_paragraph(style='List Bullet')
run = p.add_run('Function Calling：')
run.bold = True
p.add_run('AI 可以直接调用诊断工具，自动扫描分析')

doc.add_heading('3. 自动修复', level=2)
doc.add_paragraph('检测到问题后，提供修复建议', style='List Bullet')
doc.add_paragraph('支持一键自动修复（执行前确认）', style='List Bullet')

doc.add_heading('4. 双界面', level=2)
doc.add_paragraph('CLI 命令行：适合远程 SSH、脚本集成', style='List Bullet')
doc.add_paragraph('Web 仪表盘：浏览器打开，可视化展示', style='List Bullet')

# 技术栈
doc.add_heading('🛠️ 技术栈', level=1)

table2 = doc.add_table(rows=5, cols=2)
table2.style = 'Light Grid Accent 1'
table2.alignment = WD_TABLE_ALIGNMENT.CENTER

headers2 = ['组件', '技术']
for i, header in enumerate(headers2):
    table2.rows[0].cells[i].text = header
    for paragraph in table2.rows[0].cells[i].paragraphs:
        for run in paragraph.runs:
            run.bold = True

tech_data = [
    ('语言', 'Rust（高性能、内存安全）'),
    ('Web 框架', 'Axum'),
    ('前端', '单文件 HTML（零依赖，兼容旧浏览器）'),
    ('AI 集成', 'Ollama / OpenAI 兼容 API / Anthropic Claude'),
]

for i, (comp, tech) in enumerate(tech_data, 1):
    table2.rows[i].cells[0].text = comp
    table2.rows[i].cells[1].text = tech

doc.add_paragraph()

# 当前进展
doc.add_heading('📦 当前进展', level=1)
doc.add_paragraph('✅ v0.3.1 已发布', style='List Bullet')
doc.add_paragraph('✅ 支持 amd64/arm64 架构', style='List Bullet')
doc.add_paragraph('✅ 提供 deb 安装包，一键安装', style='List Bullet')
doc.add_paragraph('✅ GitHub 开源：https://github.com/fanwenzhu/kylin-doctor', style='List Bullet')

# 意见收集
doc.add_heading('💬 想听听大家的意见', level=1)
doc.add_paragraph('功能方向：你觉得还缺什么功能？（比如：远程批量诊断？定时巡检？告警通知？）', style='List Number')
doc.add_paragraph('使用场景：你在实际工作中会怎么用这类工具？', style='List Number')
doc.add_paragraph('AI 集成：对于接入 AI 做智能分析，有什么想法或顾虑？', style='List Number')
doc.add_paragraph('其他建议：任何想法都欢迎！', style='List Number')

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('欢迎大家提 issue、PR，或者直接群里讨论 🙏')
run.italic = True
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

# 保存
output_path = '/home/agent/projects/kylin-doctor/kylin-doctor-项目介绍.docx'
doc.save(output_path)
print(f'✅ 文档已保存到: {output_path}')
